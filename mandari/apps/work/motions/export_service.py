# SPDX-License-Identifier: AGPL-3.0-or-later
"""
Export Service for Motions.

Handles generation of PDF and DOCX documents from motion content,
including support for letterheads (PDF backgrounds) and proper formatting.
TipTap editor HTML is cleaned before export (comment marks stripped etc.).
"""

import io
import logging
import re
from typing import TYPE_CHECKING

from django.template.loader import render_to_string
from django.utils import timezone
from django.utils.html import escape

if TYPE_CHECKING:
    from .models import Motion, OrganizationLetterhead

logger = logging.getLogger(__name__)

#: Unterstützte Platzhalter im Signaturblock (und in Inhaltsvorlagen).
PLACEHOLDER_KEYS = ("organisation", "datum", "verantwortlich")


def build_placeholder_values(organization, responsible_name: str = "") -> dict:
    """Platzhalter-Werte für Vorlagen/Signaturblöcke einer Organisation."""
    return {
        "organisation": organization.name,
        "datum": timezone.localdate().strftime("%d.%m.%Y"),
        "verantwortlich": responsible_name,
    }


def apply_placeholders(text: str, values: dict) -> str:
    """Ersetzt {{ platzhalter }} in einem Text durch die übergebenen Werte."""
    if not text:
        return ""
    for key, value in values.items():
        text = re.sub(r"\{\{\s*" + key + r"\s*\}\}", lambda _m, v=value: v, text)
    return text


def replace_placeholders(text: str, motion: "Motion") -> str:
    """
    Ersetzt einfache Platzhalter ({{ organisation }}, {{ datum }},
    {{ verantwortlich }}) in Vorlagen-/Signaturtexten eines Dokuments.
    """
    if not text:
        return ""

    responsible = motion.responsible or motion.author
    responsible_name = ""
    if responsible:
        responsible_name = responsible.user.get_display_name()

    return apply_placeholders(text, build_placeholder_values(motion.organization, responsible_name))


def get_signature_text(motion: "Motion") -> str:
    """Signaturblock der Vorlage mit ersetzten Platzhaltern (Plaintext)."""
    if not motion.template or not motion.template.signature_block:
        return ""
    return replace_placeholders(motion.template.signature_block, motion)


def generated_letterhead_context(letterhead: "OrganizationLetterhead") -> dict:
    """
    Kontext für die HTML-Darstellung eines generierten Briefkopfs
    (Formular-Vorschau, Editor-Vorschau und PDF-Export nutzen dieselben Daten).
    """
    org = letterhead.organization
    logo = org.effective_logo if letterhead.header_logo_enabled else None

    logo_url = ""
    logo_path = ""
    if logo:
        try:
            logo_url = logo.url
            logo_path = logo.path
        except (ValueError, NotImplementedError):  # Kein File / kein FS-Storage
            logo_url = ""
            logo_path = ""

    return {
        "letterhead": letterhead,
        "org_name": org.name,
        "primary_color": org.effective_primary_color,
        "accent_enabled": letterhead.accent_color_enabled,
        "logo_url": logo_url,
        "logo_path": logo_path,
        "sender_line": letterhead.sender_line,
        "address_lines": [line for line in letterhead.address_block.splitlines() if line.strip()],
        "footer_lines": [line for line in letterhead.footer_text.splitlines() if line.strip()],
    }


def clean_editor_html(html: str) -> str:
    """
    Clean TipTap editor HTML for export.

    Strips comment marks, data attributes, and other editor-specific
    elements that shouldn't appear in exported documents.
    """
    if not html:
        return ""

    # Remove comment mark wrappers (mark/span) and keep inner text
    html = re.sub(r"<(?:mark|span)[^>]*data-comment-id[^>]*>(.*?)</(?:mark|span)>", r"\1", html, flags=re.DOTALL)

    # Remove any remaining data- attributes from all elements
    html = re.sub(r'\s+data-[a-z-]+="[^"]*"', "", html)

    # Remove editor-specific CSS classes
    html = re.sub(r'\s+class="comment-mark[^"]*"', "", html)
    html = re.sub(r'\s+class="is-editor-empty"', "", html)

    return html


class MotionExportService:
    """Service for exporting motions to various formats."""

    # Default margins in mm
    MARGIN_TOP = 30
    MARGIN_BOTTOM = 25
    MARGIN_LEFT = 25
    MARGIN_RIGHT = 25

    def export_to_pdf(self, motion: "Motion") -> bytes:
        """
        Generate a PDF from the motion content.

        Args:
            motion: The motion to export

        Returns:
            bytes: The PDF file content
        """
        # Generate HTML content
        html_content = self._render_html(motion)

        # Convert HTML to PDF
        pdf_content = self._html_to_pdf(html_content)

        # Apply PDF letterhead overlay (kind=pdf only)
        letterhead = motion.letterhead
        if letterhead and not letterhead.is_generated and letterhead.pdf_file:
            pdf_content = self._apply_letterhead(pdf_content, letterhead)

        return pdf_content

    def _pdf_margins(self, letterhead) -> dict:
        """Ränder (mm) für den PDF-Export: aus dem Briefkopf oder Defaults."""
        if letterhead:
            return {
                "top": letterhead.content_margin_top,
                "right": letterhead.content_margin_right,
                "bottom": letterhead.content_margin_bottom,
                "left": letterhead.content_margin_left,
            }
        return {
            "top": self.MARGIN_TOP,
            "right": self.MARGIN_RIGHT,
            "bottom": self.MARGIN_BOTTOM,
            "left": self.MARGIN_LEFT,
        }

    def _render_html(self, motion: "Motion") -> str:
        """Render the motion content as HTML for PDF generation."""
        # Clean editor-specific HTML before rendering
        content = clean_editor_html(motion.content or "")

        letterhead = motion.letterhead
        margins = self._pdf_margins(letterhead)
        # A4: 210 x 297 mm — Frame-Geometrie für xhtml2pdf vorberechnen
        frame = {
            "content_left": margins["left"],
            "content_top": margins["top"],
            "content_width": 210 - margins["left"] - margins["right"],
            "content_height": 297 - margins["top"] - margins["bottom"],
            "footer_top": 297 - margins["bottom"] + 3,
            "footer_height": max(margins["bottom"] - 6, 8),
        }

        signature_text = get_signature_text(motion)

        context = {
            "title": motion.title,
            "content": content,
            "margins": margins,
            "frame": frame,
            "font_family": letterhead.font_family if letterhead else "Helvetica",
            "font_size": letterhead.font_size if letterhead else 11,
            "signature_html": escape(signature_text).replace("\n", "<br>") if signature_text else "",
            "generated": bool(letterhead and letterhead.is_generated),
        }

        if letterhead and letterhead.is_generated:
            context.update(generated_letterhead_context(letterhead))

        return render_to_string("work/motions/export/pdf_template.html", context)

    def _html_to_pdf(self, html_content: str) -> bytes:
        """Convert HTML to PDF using xhtml2pdf."""
        try:
            from xhtml2pdf import pisa

            result = io.BytesIO()

            # Create PDF
            pisa_status = pisa.CreatePDF(
                src=html_content,
                dest=result,
                encoding="UTF-8",
            )

            if pisa_status.err:
                raise Exception(f"PDF generation error: {pisa_status.err}")

            return result.getvalue()

        except ImportError:
            # Fallback: Generate a simple PDF with reportlab
            return self._simple_pdf(html_content)

    def _simple_pdf(self, html_content: str) -> bytes:
        """Fallback simple PDF generation using reportlab."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=self.MARGIN_LEFT * mm,
            rightMargin=self.MARGIN_RIGHT * mm,
            topMargin=self.MARGIN_TOP * mm,
            bottomMargin=self.MARGIN_BOTTOM * mm,
        )

        styles = getSampleStyleSheet()

        # Strip HTML tags for simple text
        text = re.sub(r"<[^>]+>", "", html_content)
        text = text.replace("&nbsp;", " ").replace("&amp;", "&")

        # Add content
        story = []
        for line in text.split("\n"):
            line = line.strip()
            if line:
                story.append(Paragraph(line, styles["Normal"]))
                story.append(Spacer(1, 6))

        doc.build(story)
        return buffer.getvalue()

    def _apply_letterhead(self, pdf_content: bytes, letterhead) -> bytes:
        """
        Overlay the PDF content on the letterhead background.

        Args:
            pdf_content: The generated PDF content
            letterhead: The OrganizationLetterhead instance

        Returns:
            bytes: The merged PDF with letterhead
        """
        try:
            from copy import copy

            from pypdf import PdfReader, PdfWriter

            # Read the content PDF
            content_pdf = PdfReader(io.BytesIO(pdf_content))

            # Read the letterhead PDF
            letterhead_pdf = PdfReader(letterhead.pdf_file)
            letterhead_page = letterhead_pdf.pages[0]

            # Create output PDF
            output = PdfWriter()

            # Merge each page with the letterhead
            for content_page in content_pdf.pages:
                new_page = copy(letterhead_page)
                new_page.merge_page(content_page)
                output.add_page(new_page)

            # Write to bytes
            result = io.BytesIO()
            output.write(result)
            return result.getvalue()

        except Exception as e:
            logger.warning(f"Letterhead merge error: {e}")
            return pdf_content

    def export_to_docx(self, motion: "Motion") -> bytes:
        """
        Generate a DOCX from the motion content.

        Converts TipTap HTML to structured DOCX using python-docx.
        Letterhead data (margins, header with logo/address, footer with
        page numbers) is applied from the motion's letterhead.

        Args:
            motion: The motion to export

        Returns:
            bytes: The DOCX file content
        """
        from docx import Document
        from docx.shared import Mm, Pt

        doc = Document()
        letterhead = motion.letterhead

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = letterhead.font_family if letterhead else "Arial"
        font.size = Pt(letterhead.font_size if letterhead else 11)

        # Page margins from letterhead mm fields (or defaults)
        margins = self._pdf_margins(letterhead)
        section = doc.sections[0]
        section.top_margin = Mm(margins["top"])
        section.right_margin = Mm(margins["right"])
        section.bottom_margin = Mm(margins["bottom"])
        section.left_margin = Mm(margins["left"])

        if letterhead:
            self._docx_apply_header(doc, section, letterhead)
            self._docx_apply_footer(section, letterhead)

        # Parse and convert HTML content directly (no title/metadata header)
        content = clean_editor_html(motion.content or "")
        self._html_to_docx(doc, content)

        # Signature block from template (placeholders replaced)
        signature_text = get_signature_text(motion)
        if signature_text:
            doc.add_paragraph()
            for line in signature_text.splitlines():
                doc.add_paragraph(line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _docx_apply_header(self, doc, section, letterhead) -> None:
        """
        Kopfzeile: Logo (falls vorhanden/aktiviert) + Absenderblock rechts,
        optional Akzentlinie und Absenderzeile. Gilt für beide Briefkopf-Arten;
        ohne gepflegte generated-Felder wird ein schlichter Kopf mit dem
        Organisationsnamen in der Primärfarbe gesetzt.
        """
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt

        org = letterhead.organization
        ctx = generated_letterhead_context(letterhead)
        primary_rgb = self._hex_to_rgbcolor(ctx["primary_color"])

        header = section.header
        header.is_linked_to_previous = False

        has_generated_fields = bool(ctx["logo_path"] or ctx["address_lines"] or ctx["sender_line"])

        if not has_generated_fields:
            # Schlichter Kopf: Organisationsname in Primärfarbe
            para = header.paragraphs[0]
            run = para.add_run(org.name)
            run.font.bold = True
            run.font.size = Pt(14)
            if primary_rgb is not None:
                run.font.color.rgb = primary_rgb
            return

        # Logo + Absenderblock nebeneinander (Tabelle mit 2 Spalten)
        usable_width = section.page_width - section.left_margin - section.right_margin
        table = header.add_table(rows=1, cols=2, width=usable_width)
        cells = table.rows[0].cells
        cells[0].width = usable_width // 2
        cells[1].width = usable_width - usable_width // 2

        logo_cell = cells[0].paragraphs[0]
        if ctx["logo_path"]:
            try:
                logo_run = logo_cell.add_run()
                logo_run.add_picture(ctx["logo_path"], height=Mm(15))
            except Exception as e:
                logger.warning(f"DOCX-Logo konnte nicht eingefügt werden: {e}")

        address_para = cells[1].paragraphs[0]
        address_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for index, line in enumerate(ctx["address_lines"]):
            if index:
                address_para.add_run().add_break()
            run = address_para.add_run(line)
            run.font.size = Pt(9)

        # Akzentlinie in Primärfarbe (Rahmen unter einem Leerabsatz)
        if letterhead.accent_color_enabled:
            accent_para = header.add_paragraph()
            self._docx_set_bottom_border(accent_para, ctx["primary_color"])

        # Absenderzeile klein
        if ctx["sender_line"]:
            sender_para = header.add_paragraph()
            run = sender_para.add_run(ctx["sender_line"])
            run.font.size = Pt(7)
            run.font.underline = True

    def _docx_apply_footer(self, section, letterhead) -> None:
        """Fußzeile: footer_text-Zeilen + Seitenzahl-Feld."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        ctx = generated_letterhead_context(letterhead)

        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for index, line in enumerate(ctx["footer_lines"]):
            if index:
                para.add_run().add_break()
            run = para.add_run(line)
            run.font.size = Pt(8)

        page_para = footer.add_paragraph()
        page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prefix_run = page_para.add_run("Seite ")
        prefix_run.font.size = Pt(8)
        self._docx_add_page_number_field(page_para)

    @staticmethod
    def _docx_add_page_number_field(paragraph) -> None:
        """Fügt ein PAGE-Feld (dynamische Seitenzahl) in den Absatz ein."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1"
        run.append(text)
        fld.append(run)
        paragraph._p.append(fld)

    @staticmethod
    def _docx_set_bottom_border(paragraph, hex_color: str) -> None:
        """Setzt eine untere Rahmenlinie (Akzentlinie) am Absatz."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        p_pr = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), hex_color.lstrip("#"))
        borders.append(bottom)
        p_pr.append(borders)

    @staticmethod
    def _hex_to_rgbcolor(hex_color: str):
        """#RRGGBB → docx RGBColor (None bei ungültigem Wert)."""
        from docx.shared import RGBColor

        value = (hex_color or "").lstrip("#")
        if len(value) != 6:
            return None
        try:
            return RGBColor.from_string(value.upper())
        except ValueError:
            return None

    def _html_to_docx(self, doc, html: str):
        """
        Convert cleaned HTML to DOCX paragraphs in document order.

        Processes elements sequentially to maintain the original document
        structure rather than grouping by type.
        """
        from html.parser import HTMLParser

        from docx.shared import Pt

        if not html:
            return

        class DocxHTMLConverter(HTMLParser):
            """Sequential HTML-to-DOCX converter."""

            def __init__(self, doc):
                super().__init__()
                self.doc = doc
                self._text_buf: list[str] = []
                self._tag_stack: list[str] = []
                self._list_type_stack: list[str] = []  # 'ul' or 'ol'
                self._in_heading = 0  # heading level (1-3) or 0
                self._in_blockquote = False
                self._bold = False
                self._italic = False
                self._underline = False
                self._strike = False
                self._indent_px = 0

            def _flush_text(self):
                """Return accumulated text and reset buffer."""
                text = "".join(self._text_buf).strip()
                self._text_buf.clear()
                return text

            def handle_starttag(self, tag, attrs):
                attrs_dict = dict(attrs)
                tag_lower = tag.lower()
                self._tag_stack.append(tag_lower)

                if tag_lower in ("h1", "h2", "h3"):
                    self._flush_text()
                    self._in_heading = int(tag_lower[1])
                elif tag_lower == "blockquote":
                    self._flush_text()
                    self._in_blockquote = True
                elif tag_lower == "ul":
                    self._list_type_stack.append("ul")
                elif tag_lower == "ol":
                    self._list_type_stack.append("ol")
                elif tag_lower == "li":
                    self._flush_text()
                elif tag_lower == "p":
                    self._flush_text()
                    # Check for indent via margin-left style
                    style = attrs_dict.get("style", "")
                    self._indent_px = 0
                    if "margin-left" in style:
                        match = re.search(r"margin-left:\s*(\d+)", style)
                        if match:
                            self._indent_px = int(match.group(1))
                elif tag_lower == "br":
                    self._text_buf.append("\n")
                elif tag_lower == "hr":
                    self._flush_text()
                    # Add a thin horizontal line as paragraph
                    para = self.doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(8)
                    para.paragraph_format.space_after = Pt(8)
                    run = para.add_run("_" * 60)
                    run.font.color.rgb = None  # default color
                    run.font.size = Pt(6)
                elif tag_lower == "strong" or tag_lower == "b":
                    self._bold = True
                elif tag_lower == "em" or tag_lower == "i":
                    self._italic = True
                elif tag_lower == "u":
                    self._underline = True
                elif tag_lower in ("s", "del", "strike"):
                    self._strike = True

            def handle_endtag(self, tag):
                tag_lower = tag.lower()

                if tag_lower in ("h1", "h2", "h3"):
                    text = self._flush_text()
                    if text:
                        self.doc.add_heading(text, level=self._in_heading)
                    self._in_heading = 0
                elif tag_lower == "blockquote":
                    text = self._flush_text()
                    if text:
                        para = self.doc.add_paragraph(text)
                        para.paragraph_format.left_indent = Pt(36)
                        for run in para.runs:
                            run.font.italic = True
                    self._in_blockquote = False
                elif tag_lower in ("ul", "ol"):
                    if self._list_type_stack:
                        self._list_type_stack.pop()
                elif tag_lower == "li":
                    text = self._flush_text()
                    if text:
                        list_type = self._list_type_stack[-1] if self._list_type_stack else "ul"
                        style = "List Number" if list_type == "ol" else "List Bullet"
                        para = self.doc.add_paragraph(text, style=style)
                        # Indent nested lists
                        depth = len(self._list_type_stack)
                        if depth > 1:
                            para.paragraph_format.left_indent = Pt(18 * (depth - 1))
                elif tag_lower == "p":
                    text = self._flush_text()
                    if text:
                        para = self.doc.add_paragraph(text)
                        if self._indent_px > 0:
                            para.paragraph_format.left_indent = Pt(self._indent_px * 0.75)
                    self._indent_px = 0
                elif tag_lower == "strong" or tag_lower == "b":
                    self._bold = False
                elif tag_lower == "em" or tag_lower == "i":
                    self._italic = False
                elif tag_lower == "u":
                    self._underline = False
                elif tag_lower in ("s", "del", "strike"):
                    self._strike = False

                if self._tag_stack and self._tag_stack[-1] == tag_lower:
                    self._tag_stack.pop()

            def handle_data(self, data):
                self._text_buf.append(data)

        converter = DocxHTMLConverter(doc)
        converter.feed(html)


# Singleton instance
motion_export_service = MotionExportService()

# SPDX-License-Identifier: AGPL-3.0-or-later
"""
Import Service for Motions/Documents.

Imports PDF and DOCX files as documents, extracting text content
and storing the original file as an attachment.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from typing import TYPE_CHECKING

from django.core.files.uploadedfile import UploadedFile

from insight_core.services.document_extraction import extract_text_from_file

if TYPE_CHECKING:
    from apps.tenants.models import Membership, Organization

    from .models import Motion, MotionDocument, MotionType


logger = logging.getLogger(__name__)


@dataclass
class ImportResult:
    """Result of a file import operation (PDF or DOCX)."""

    success: bool
    motion: Motion | None = None
    document: MotionDocument | None = None
    error: str | None = None
    extracted_text_length: int = 0
    ocr_performed: bool = False


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


class MotionImportService:
    """
    Service for importing files (PDF/DOCX) as Motion documents.

    Extracts text content and creates Motion instances
    with the extracted text as editor-ready HTML content.
    """

    @classmethod
    def import_pdf(
        cls,
        pdf_file: UploadedFile,
        organization: Organization,
        author: Membership,
        motion_type: MotionType | None = None,
        title: str | None = None,
        visibility: str = "private",
    ) -> ImportResult:
        """
        Import a PDF file as a new Motion document.

        Args:
            pdf_file: The uploaded PDF file
            organization: The organization to create the motion in
            author: The membership creating the motion
            motion_type: Optional document type
            title: Optional title (defaults to filename)
            visibility: Visibility setting (default: private)

        Returns:
            ImportResult with the created motion and document
        """
        from .models import Motion, MotionDocument

        try:
            # Read file content
            file_content = pdf_file.read()
            pdf_file.seek(0)  # Reset for later save

            # Extract text
            text_content, ocr_performed, page_count = extract_text_from_file(
                data=file_content,
                mime_type="application/pdf",
                file_name=pdf_file.name,
            )

            # Generate title from filename if not provided
            if not title:
                # Remove .pdf extension and clean up
                title = pdf_file.name
                if title.lower().endswith(".pdf"):
                    title = title[:-4]
                # Clean up common patterns
                title = title.replace("_", " ").replace("-", " ").strip()
                # Capitalize first letter
                if title:
                    title = title[0].upper() + title[1:]

            # Create the Motion
            motion = Motion(
                organization=organization,
                author=author,
                title=title,
                status="draft",
                visibility=visibility,
            )

            # Set document type if provided
            if motion_type:
                motion.document_type = motion_type

            # Set content (encrypted)
            if text_content:
                # Wrap in basic HTML if it's plain text
                if not text_content.strip().startswith("<"):
                    # Convert line breaks to paragraphs
                    paragraphs = text_content.split("\n\n")
                    html_content = "\n".join(f"<p>{p.strip()}</p>" for p in paragraphs if p.strip())
                    motion.set_content_encrypted(html_content)
                else:
                    motion.set_content_encrypted(text_content)
            else:
                motion.set_content_encrypted(
                    "<p><em>Text konnte nicht extrahiert werden. Bitte überprüfen Sie das Original-PDF.</em></p>"
                )

            motion.save()

            # Create MotionDocument attachment
            document = MotionDocument(
                motion=motion,
                file=pdf_file,
                filename=pdf_file.name,
                mime_type="application/pdf",
                file_size=len(file_content),
                text_content=text_content[:50000] if text_content else "",  # Limit for search
                uploaded_by=author,
            )
            document.save()

            logger.info(
                f"Imported PDF '{pdf_file.name}' as motion '{motion.title}' "
                f"(ID: {motion.id}, text length: {len(text_content)}, OCR: {ocr_performed})"
            )

            return ImportResult(
                success=True,
                motion=motion,
                document=document,
                extracted_text_length=len(text_content),
                ocr_performed=ocr_performed,
            )

        except Exception as e:
            logger.exception(f"Failed to import PDF '{pdf_file.name}': {e}")
            return ImportResult(
                success=False,
                error=str(e),
            )

    @classmethod
    def import_docx(
        cls,
        docx_file: UploadedFile,
        organization: Organization,
        author: Membership,
        motion_type: MotionType | None = None,
        title: str | None = None,
        visibility: str = "private",
    ) -> ImportResult:
        """
        Import a DOCX file as a new Motion document.

        Extracts text and basic formatting from DOCX using python-docx,
        converts to HTML for the TipTap editor.

        Args:
            docx_file: The uploaded DOCX file
            organization: The organization to create the motion in
            author: The membership creating the motion
            motion_type: Optional document type
            title: Optional title (defaults to filename)
            visibility: Visibility setting (default: private)

        Returns:
            ImportResult with the created motion and document
        """
        from .models import Motion, MotionDocument

        try:
            from docx import Document

            file_content = docx_file.read()
            docx_file.seek(0)

            # Parse the DOCX
            import io

            doc = Document(io.BytesIO(file_content))

            # Convert to HTML
            html_parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect heading styles
                style_name = (para.style.name or "").lower()
                if style_name.startswith("heading 1") or style_name == "title":
                    html_parts.append(f"<h1>{_escape_html(text)}</h1>")
                elif style_name.startswith("heading 2"):
                    html_parts.append(f"<h2>{_escape_html(text)}</h2>")
                elif style_name.startswith("heading 3"):
                    html_parts.append(f"<h3>{_escape_html(text)}</h3>")
                elif style_name.startswith("list"):
                    html_parts.append(f"<li>{_escape_html(text)}</li>")
                elif style_name == "quote" or style_name.startswith("block"):
                    html_parts.append(f"<blockquote><p>{_escape_html(text)}</p></blockquote>")
                else:
                    # Build inline formatting
                    inline_html = cls._runs_to_html(para.runs)
                    html_parts.append(f"<p>{inline_html}</p>")

            # Wrap adjacent list items in <ul>
            html_content = "\n".join(html_parts)
            html_content = re.sub(
                r"((?:<li>.*?</li>\n?)+)",
                r"<ul>\1</ul>",
                html_content,
            )

            # Generate title from filename if not provided
            if not title:
                title = docx_file.name
                for ext in (".docx", ".DOCX", ".doc", ".DOC"):
                    if title.endswith(ext):
                        title = title[: -len(ext)]
                        break
                title = title.replace("_", " ").replace("-", " ").strip()
                if title:
                    title = title[0].upper() + title[1:]

            # Create the Motion
            motion = Motion(
                organization=organization,
                author=author,
                title=title,
                status="draft",
                visibility=visibility,
            )
            if motion_type:
                motion.document_type = motion_type

            if html_content.strip():
                motion.set_content_encrypted(html_content)
            else:
                motion.set_content_encrypted(
                    "<p><em>Kein Text im Dokument gefunden.</em></p>"
                )
            motion.save()

            # Create attachment
            document = MotionDocument(
                motion=motion,
                file=docx_file,
                filename=docx_file.name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_size=len(file_content),
                text_content=re.sub(r"<[^>]+>", "", html_content)[:50000],
                uploaded_by=author,
            )
            document.save()

            logger.info(
                f"Imported DOCX '{docx_file.name}' as motion '{motion.title}' "
                f"(ID: {motion.id}, text length: {len(html_content)})"
            )

            return ImportResult(
                success=True,
                motion=motion,
                document=document,
                extracted_text_length=len(html_content),
            )

        except ImportError:
            logger.error("python-docx is not installed. Install with: pip install python-docx")
            return ImportResult(success=False, error="DOCX-Import nicht verfügbar (python-docx fehlt)")
        except Exception as e:
            logger.exception(f"Failed to import DOCX '{docx_file.name}': {e}")
            return ImportResult(success=False, error=str(e))

    @staticmethod
    def _runs_to_html(runs) -> str:
        """Convert paragraph runs to HTML with inline formatting."""
        if not runs:
            return ""
        parts = []
        for run in runs:
            text = _escape_html(run.text)
            if not text:
                continue
            if run.bold:
                text = f"<strong>{text}</strong>"
            if run.italic:
                text = f"<em>{text}</em>"
            if run.underline:
                text = f"<u>{text}</u>"
            if run.font and run.font.strike:
                text = f"<s>{text}</s>"
            parts.append(text)
        return "".join(parts)

    @classmethod
    def import_file(
        cls,
        uploaded_file: UploadedFile,
        organization: Organization,
        author: Membership,
        motion_type: MotionType | None = None,
        title: str | None = None,
        visibility: str = "private",
    ) -> ImportResult:
        """
        Import a file (PDF or DOCX) as a new Motion document.

        Dispatches to the appropriate import method based on file type.
        """
        name = (uploaded_file.name or "").lower()
        if name.endswith(".docx"):
            return cls.import_docx(
                docx_file=uploaded_file,
                organization=organization,
                author=author,
                motion_type=motion_type,
                title=title,
                visibility=visibility,
            )
        else:
            return cls.import_pdf(
                pdf_file=uploaded_file,
                organization=organization,
                author=author,
                motion_type=motion_type,
                title=title,
                visibility=visibility,
            )

    @classmethod
    def import_multiple_files(
        cls,
        files: list[UploadedFile],
        organization: Organization,
        author: Membership,
        motion_type: MotionType | None = None,
        visibility: str = "private",
    ) -> list[ImportResult]:
        """Import multiple files (PDF/DOCX) as Motion documents."""
        results = []
        for f in files:
            result = cls.import_file(
                uploaded_file=f,
                organization=organization,
                author=author,
                motion_type=motion_type,
                visibility=visibility,
            )
            results.append(result)
        return results

    @classmethod
    def import_multiple_pdfs(
        cls,
        pdf_files: list[UploadedFile],
        organization: Organization,
        author: Membership,
        motion_type: MotionType | None = None,
        visibility: str = "private",
    ) -> list[ImportResult]:
        """
        Import multiple PDF files as Motion documents.

        Args:
            pdf_files: List of uploaded PDF files
            organization: The organization to create motions in
            author: The membership creating the motions
            motion_type: Optional document type for all imports
            visibility: Visibility setting for all imports

        Returns:
            List of ImportResult objects
        """
        results = []
        for pdf_file in pdf_files:
            result = cls.import_pdf(
                pdf_file=pdf_file,
                organization=organization,
                author=author,
                motion_type=motion_type,
                visibility=visibility,
            )
            results.append(result)
        return results


# Singleton instance for convenience
motion_import_service = MotionImportService()

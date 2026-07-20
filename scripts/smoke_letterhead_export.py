# SPDX-License-Identifier: AGPL-3.0-or-later
"""
Smoke-Test: Briefköpfe aus dem Corporate Design + Export-Integration (Etappe 2).

Läuft gegen eine frische SQLite-Instanz mit django.test.Client:
    python scripts/smoke_letterhead_export.py

Prüft:
- Generierter Briefkopf (CRUD, Defaults aus Org-Daten)
- PDF-Export mit generiertem Briefkopf (%PDF, Absenderzeile, Fußzeile, Titel, Signatur)
- PDF-Export mit PDF-Briefkopf (Overlay-Pfad unverändert)
- DOCX-Export (Header/Footer-Texte, Logo-Run, Seitenzahl-Feld, Ränder)
- Editor-Preview-Endpoint + Formular-Live-Vorschau
- Vorlagen: Default-Exklusivität je Typ, Platzhalter-Ersetzung, Vorschau
- Organisations-Grenzen (fremde Org -> 404 auf Preview/CRUD)
"""

import base64
import io
import os
import secrets
import sys
import tempfile
from pathlib import Path

# --- Umgebung VOR django.setup() konfigurieren -------------------------------
PROJECT_DIR = Path(__file__).resolve().parent.parent / "mandari"
sys.path.insert(0, str(PROJECT_DIR))

_tmp_dir = Path(tempfile.mkdtemp(prefix="mandari_smoke_lh_"))
_db_path = _tmp_dir / "smoke.sqlite3"
os.environ["DJANGO_SETTINGS_MODULE"] = "mandari.settings"
os.environ["DEBUG"] = "true"
os.environ["DATABASE_URL"] = f"sqlite:///{_db_path.as_posix()}"
os.environ["ENCRYPTION_MASTER_KEY"] = base64.b64encode(secrets.token_bytes(32)).decode()
os.environ["ELASTICSEARCH_AUTO_INDEX"] = "False"
os.environ["MANDARI_SYNC_WATCHDOG"] = "0"  # DB-Schreiber-Thread stoert SQLite-Migrationen
os.environ["EMAIL_BACKEND"] = "django.core.mail.backends.locmem.EmailBackend"
os.environ["ALLOWED_HOSTS"] = "testserver,localhost"

import django  # noqa: E402

django.setup()

from django.conf import settings  # noqa: E402

settings.MEDIA_ROOT = str(_tmp_dir / "media")

from django.core.files.base import ContentFile  # noqa: E402
from django.core.files.uploadedfile import SimpleUploadedFile  # noqa: E402
from django.core.management import call_command  # noqa: E402
from django.test import Client  # noqa: E402
from django.test.utils import setup_test_environment  # noqa: E402
from django.utils import timezone  # noqa: E402

setup_test_environment()
call_command("migrate", verbosity=0, interactive=False)

from apps.accounts.models import User  # noqa: E402
from apps.tenants.models import Membership, Organization, Role  # noqa: E402
from apps.work.motions.export_service import replace_placeholders  # noqa: E402
from apps.work.motions.models import (  # noqa: E402
    Motion,
    MotionTemplate,
    MotionType,
    OrganizationLetterhead,
)

PASS = 0
FAIL = 0

# 1x1-PNG (rot) für das Organisations-Logo
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8z8BQDwAEhQGAhKmMIQAAAABJRU5ErkJggg=="
)


def check(name, condition, detail=""):
    global PASS, FAIL
    if condition:
        PASS += 1
        print(f"  OK   {name}")
    else:
        FAIL += 1
        print(f"  FAIL {name} {detail}")


def make_org(name, slug):
    org = Organization.objects.create(
        name=name,
        slug=slug,
        address="Rathausplatz 1\n12345 Musterstadt",
        contact_email=f"kontakt@{slug}.example.org",
        contact_phone="0123 456789",
        website="https://beispiel.example.org",
        primary_color="#146a3a",
    )
    role = Role.objects.filter(organization=org, is_admin=True).first()
    if role is None:
        role = Role.objects.create(organization=org, name="Administrator", is_admin=True)
    return org, role


def make_member(org, role, email):
    user = User.objects.create_user(email=email, password="test1234!")
    membership = Membership.objects.create(user=user, organization=org)
    membership.roles.add(role)
    return user, membership


def client_for(user):
    client = Client()
    client.force_login(user)
    return client


def make_mock_pdf(text) -> bytes:
    """Erzeugt ein einfaches PDF (Mock-Briefkopf-Hintergrund)."""
    from xhtml2pdf import pisa

    buf = io.BytesIO()
    pisa.CreatePDF(src=f"<html><body><p>{text}</p></body></html>", dest=buf, encoding="UTF-8")
    return buf.getvalue()


def pdf_text(pdf_bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


print("=== Setup ===")
org_a, role_a = make_org("Fraktion A", "fraktion-a")
org_b, role_b = make_org("Fraktion B", "fraktion-b")

# Logo für Org A setzen (Corporate Design)
org_a.logo.save("logo.png", ContentFile(PNG_1PX), save=True)

user_admin, m_admin = make_member(org_a, role_a, "admin@example.org")
user_foreign, m_foreign = make_member(org_b, role_b, "fremd@example.org")

c_admin = client_for(user_admin)
c_foreign = client_for(user_foreign)

DOCS = f"/work/{org_a.slug}/documents"
SETTINGS = f"/work/{org_a.slug}/organization/documents"

doc_type = MotionType.objects.create(organization=org_a, name="Antrag", slug="antrag")
other_type = MotionType.objects.create(organization=org_a, name="Anfrage", slug="anfrage")

# --- 1. Generierten Briefkopf anlegen (CRUD) ---------------------------------
print("=== 1. Generierter Briefkopf: CRUD + Defaults ===")

resp = c_admin.get(f"{SETTINGS}/letterheads/create/")
check("Create-Formular 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
check("Defaults aus Org-Daten (Name)", "Fraktion A" in content)
check("Defaults aus Org-Daten (Adresse)", "Rathausplatz 1" in content)

resp = c_admin.post(
    f"{SETTINGS}/letterheads/create/",
    {
        "name": "CD-Briefkopf",
        "description": "Generiert aus dem Corporate Design",
        "kind": "generated",
        "header_logo_enabled": "on",
        "accent_color_enabled": "on",
        "sender_line": "Fraktion A · Rathausplatz 1 · 12345 Musterstadt",
        "address_block": "Fraktion A\nRathausplatz 1\n12345 Musterstadt",
        "footer_text": "kontakt@fraktion-a.example.org · Tel. 0123 456789\nwww.beispiel.example.org",
        "content_margin_top": "45",
        "content_margin_left": "25",
        "content_margin_right": "20",
        "content_margin_bottom": "30",
        "font_family": "Helvetica",
        "font_size": "11",
        "is_default": "on",
    },
)
check("Generierter Briefkopf angelegt (302)", resp.status_code == 302, f"got {resp.status_code}")
lh_generated = OrganizationLetterhead.objects.get(organization=org_a, name="CD-Briefkopf")
check("kind = generated", lh_generated.kind == "generated")
check("Kein PDF nötig", not lh_generated.pdf_file)

# PDF-Briefkopf (Overlay) anlegen
mock_pdf = make_mock_pdf("MUSTER-BRIEFKOPF-HINTERGRUND")
resp = c_admin.post(
    f"{SETTINGS}/letterheads/create/",
    {
        "name": "PDF-Briefkopf",
        "kind": "pdf",
        "pdf_file": SimpleUploadedFile("briefkopf.pdf", mock_pdf, content_type="application/pdf"),
        "content_margin_top": "60",
        "content_margin_left": "25",
        "content_margin_right": "20",
        "content_margin_bottom": "30",
        "font_family": "Arial",
        "font_size": "11",
    },
)
check("PDF-Briefkopf angelegt (302)", resp.status_code == 302, f"got {resp.status_code}")
lh_pdf = OrganizationLetterhead.objects.get(organization=org_a, name="PDF-Briefkopf")
check("kind = pdf", lh_pdf.kind == "pdf")

# kind=pdf ohne Datei wird abgelehnt
resp = c_admin.post(f"{SETTINGS}/letterheads/create/", {"name": "Kaputt", "kind": "pdf"})
created = OrganizationLetterhead.objects.filter(organization=org_a, name="Kaputt").exists()
check("PDF-Briefkopf ohne Datei abgelehnt", resp.status_code == 200 and not created)

# --- 2. Formular-Live-Vorschau + Editor-Preview ------------------------------
print("=== 2. Vorschau-Endpoints ===")

resp = c_admin.get(
    f"{SETTINGS}/letterheads/preview/",
    {
        "header_logo_enabled": "on",
        "accent_color_enabled": "on",
        "sender_line": "Fraktion A · Rathausplatz 1",
        "address_block": "Fraktion A\nRathausplatz 1",
        "footer_text": "kontakt@fraktion-a.example.org",
    },
)
check("Formular-Vorschau 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
check("Vorschau enthält Absenderzeile", "Fraktion A · Rathausplatz 1" in content)
check("Vorschau enthält Fußzeile", "kontakt@fraktion-a.example.org" in content)
check("Vorschau enthält Logo", 'src="' in content and "logo" in content)
check("Vorschau enthält Akzentfarbe", "#146a3a" in content)

resp = c_admin.get(f"{DOCS}/letterheads/{lh_generated.id}/preview/")
check("Editor-Preview 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
check("Editor-Preview enthält Absenderblock", "Rathausplatz 1" in content)
check("Editor-Preview enthält Absenderzeile", "Fraktion A · Rathausplatz 1 · 12345 Musterstadt" in content)

# --- 3. Vorlagen: Default je Typ + Platzhalter --------------------------------
print("=== 3. Vorlagen ===")

tpl_data = {
    "name": "Standardantrag",
    "description": "Basis-Vorlage",
    "motion_type": str(doc_type.id),
    "letterhead": str(lh_generated.id),
    "content_template": "<p>Antrag der {{ organisation }} vom {{ datum }}</p>",
    "signature_block": "Mit freundlichen Grüßen\n{{ verantwortlich }}\n{{ organisation }}, {{ datum }}",
    "is_default": "on",
    "is_active": "on",
}
resp = c_admin.post(f"{SETTINGS}/templates/create/", tpl_data)
check("Vorlage 1 angelegt (302)", resp.status_code == 302, f"got {resp.status_code}")
tpl1 = MotionTemplate.objects.get(organization=org_a, name="Standardantrag")

resp = c_admin.post(
    f"{SETTINGS}/templates/create/",
    {**tpl_data, "name": "Anfrage-Vorlage", "motion_type": str(other_type.id)},
)
check("Vorlage 2 (anderer Typ) angelegt", resp.status_code == 302)

resp = c_admin.post(f"{SETTINGS}/templates/create/", {**tpl_data, "name": "Zweitantrag"})
check("Vorlage 3 (gleicher Typ, Default) angelegt", resp.status_code == 302)

tpl1.refresh_from_db()
tpl2 = MotionTemplate.objects.get(organization=org_a, name="Anfrage-Vorlage")
tpl3 = MotionTemplate.objects.get(organization=org_a, name="Zweitantrag")
check("Default-Exklusivität je Typ: Vorlage 1 zurückgesetzt", tpl1.is_default is False)
check("Default anderer Typ bleibt", tpl2.is_default is True)
check("Neue Vorlage ist Default", tpl3.is_default is True)

# Vorlagen-Vorschau
resp = c_admin.get(f"{SETTINGS}/templates/{tpl1.id}/preview/")
check("Vorlagen-Vorschau 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
today = timezone.localdate().strftime("%d.%m.%Y")
check("Vorschau: Platzhalter organisation ersetzt", "Antrag der Fraktion A" in content)
check("Vorschau: Platzhalter datum ersetzt", today in content)
check("Vorschau: Briefkopf enthalten", "Rathausplatz 1" in content)

# --- 4. Dokument anlegen: Platzhalter in Inhaltsvorlage -----------------------
print("=== 4. Dokument aus Vorlage ===")

resp = c_admin.post(
    f"{DOCS}/create/",
    {
        "title": "Radweg Musterstraße",
        "document_type": str(doc_type.id),
        "template": str(tpl1.id),
        "letterhead": str(lh_generated.id),
    },
)
check("Dokument angelegt (302)", resp.status_code == 302, f"got {resp.status_code}")
motion = Motion.objects.get(organization=org_a, title="Radweg Musterstraße")
check("Briefkopf übernommen", motion.letterhead_id == lh_generated.id)
check(
    "Platzhalter in Inhaltsvorlage ersetzt",
    f"Antrag der Fraktion A vom {today}" in motion.content,
    motion.content[:120],
)

sig = replace_placeholders(motion.template.signature_block, motion)
check("Platzhalter verantwortlich ersetzt", "admin@example.org" in sig or "{{" not in sig, sig)
check("Signatur ohne Platzhalter-Reste", "{{" not in sig, sig)

# --- 5. PDF-Export mit generiertem Briefkopf ----------------------------------
print("=== 5. PDF-Export (generiert) ===")

resp = c_admin.get(f"{DOCS}/{motion.id}/export/", {"format": "pdf"})
check("PDF-Export 200", resp.status_code == 200, f"got {resp.status_code}")
pdf_bytes = resp.content
check("PDF-Magic (%PDF)", pdf_bytes.startswith(b"%PDF"))
text = pdf_text(pdf_bytes)
check("PDF enthält Dokumenttitel", "Radweg Musterstraße" in text, text[:200])
check("PDF enthält Absenderzeile", "12345 Musterstadt" in text)
check("PDF enthält Fußzeilen-Text", "kontakt@fraktion-a.example.org" in text)
check("PDF enthält Seitenzahl", "Seite 1 von 1" in text.replace("\n", " "))
check("PDF enthält Signaturblock", "Mit freundlichen Grüßen" in text)
check("PDF enthält Inhalt", "Antrag der Fraktion A" in text)

# --- 6. PDF-Export mit PDF-Briefkopf (Overlay unverändert) ---------------------
print("=== 6. PDF-Export (Overlay) ===")

motion.letterhead = lh_pdf
motion.save(update_fields=["letterhead"])

resp = c_admin.get(f"{DOCS}/{motion.id}/export/", {"format": "pdf"})
check("Overlay-PDF-Export 200", resp.status_code == 200, f"got {resp.status_code}")
pdf_bytes = resp.content
check("Overlay-PDF-Magic", pdf_bytes.startswith(b"%PDF"))
text = pdf_text(pdf_bytes)
check("Overlay: Briefkopf-Hintergrund enthalten", "MUSTER-BRIEFKOPF-HINTERGRUND" in text, text[:200])
check("Overlay: Inhalt enthalten", "Antrag der Fraktion A" in text)
check("Overlay: kein generierter Kopf", "12345 Musterstadt" not in text)

# --- 7. DOCX-Export ------------------------------------------------------------
print("=== 7. DOCX-Export ===")

motion.letterhead = lh_generated
motion.save(update_fields=["letterhead"])

resp = c_admin.get(f"{DOCS}/{motion.id}/export/", {"format": "docx"})
check("DOCX-Export 200", resp.status_code == 200, f"got {resp.status_code}")

from docx import Document  # noqa: E402

doc = Document(io.BytesIO(resp.content))
section = doc.sections[0]
# DOCX speichert Ränder in Twips — auf ganze mm gerundet vergleichen
check("DOCX Ränder aus mm-Feldern (oben)", round(section.top_margin.mm) == 45, str(section.top_margin.mm))
check("DOCX Ränder aus mm-Feldern (links)", round(section.left_margin.mm) == 25, str(section.left_margin.mm))

header_text = "\n".join(p.text for p in section.header.paragraphs)
header_tables = section.header.tables
table_text = ""
if header_tables:
    table_text = "\n".join(cell.text for row in header_tables[0].rows for cell in row.cells)
check("DOCX Header enthält Absenderblock", "Rathausplatz 1" in (header_text + table_text))
check("DOCX Header enthält Absenderzeile", "Fraktion A · Rathausplatz 1 · 12345 Musterstadt" in header_text)

header_xml = section.header._element.xml
check("DOCX Logo-Run vorhanden", "drawing" in header_xml or "pic:pic" in header_xml)

footer_text = "\n".join(p.text for p in section.footer.paragraphs)
check("DOCX Footer enthält Fußzeilen-Text", "kontakt@fraktion-a.example.org" in footer_text)
footer_xml = section.footer._element.xml
check("DOCX Footer enthält Seitenzahl-Feld", 'w:instr="PAGE"' in footer_xml)

body_text = "\n".join(p.text for p in doc.paragraphs)
check("DOCX enthält Inhalt", "Antrag der Fraktion A" in body_text)
check("DOCX enthält Signaturblock", "Mit freundlichen Grüßen" in body_text)

# DOCX mit PDF-Briefkopf ohne generated-Felder: schlichter Kopf mit Org-Name
motion.letterhead = lh_pdf
motion.save(update_fields=["letterhead"])
resp = c_admin.get(f"{DOCS}/{motion.id}/export/", {"format": "docx"})
check("DOCX-Export (PDF-Briefkopf) 200", resp.status_code == 200)
doc2 = Document(io.BytesIO(resp.content))
header2_text = "\n".join(p.text for p in doc2.sections[0].header.paragraphs)
check("DOCX schlichter Kopf mit Org-Name", "Fraktion A" in header2_text, header2_text)

motion.letterhead = lh_generated
motion.save(update_fields=["letterhead"])

# --- 8. Organisations-Grenzen ---------------------------------------------------
print("=== 8. Organisations-Grenzen ===")

FOREIGN_DOCS = f"/work/{org_b.slug}/documents"
FOREIGN_SETTINGS = f"/work/{org_b.slug}/organization/documents"

resp = c_foreign.get(f"{FOREIGN_DOCS}/letterheads/{lh_generated.id}/preview/")
check("Fremde Org: Editor-Preview 404", resp.status_code == 404, f"got {resp.status_code}")

resp = c_foreign.get(f"{FOREIGN_SETTINGS}/letterheads/{lh_generated.id}/")
check("Fremde Org: Briefkopf-Edit 404", resp.status_code == 404, f"got {resp.status_code}")

resp = c_foreign.post(f"{FOREIGN_SETTINGS}/letterheads/{lh_generated.id}/delete/")
check("Fremde Org: Briefkopf-Delete 404", resp.status_code == 404, f"got {resp.status_code}")

resp = c_foreign.get(f"{FOREIGN_SETTINGS}/templates/{tpl1.id}/preview/")
check("Fremde Org: Vorlagen-Vorschau 404", resp.status_code == 404, f"got {resp.status_code}")
lh_still_there = OrganizationLetterhead.objects.filter(id=lh_generated.id).exists()
check("Briefkopf nicht gelöscht", lh_still_there)

# --- 9. Branding-Kachel + Editor-Kontext ----------------------------------------
print("=== 9. Einstellungen + Editor ===")

resp = c_admin.get(f"{SETTINGS}/")
check("Dokument-Einstellungen 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
check("Branding-Kachel: Corporate Design", "Corporate Design" in content)
check("Branding-Kachel: Briefkopf-Status generiert", "Briefkopf: generiert" in content)
check("Branding-Kachel: Primärfarbe", "#146a3a" in content)
check("Branding-Kachel: Link zu Org-Einstellungen", f"/work/{org_a.slug}/organization/" in content)

resp = c_admin.get(f"{DOCS}/{motion.id}/")
check("Editor 200", resp.status_code == 200, f"got {resp.status_code}")
content = resp.content.decode()
check("Editor: generierter Briefkopf im JSON", '"kind": "generated"' in content)
check("Editor: Preview-URL im JSON", f"/documents/letterheads/{lh_generated.id}/preview/" in content)

# --- Ergebnis -------------------------------------------------------------------
print()
print(f"=== Ergebnis: {PASS} OK, {FAIL} FAIL ===")
sys.exit(1 if FAIL else 0)
